import io
import logging
import os
import re
from typing import List

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field


# ============================================================
# CareerPilot AI
# Production-ready FastAPI backend
# ============================================================

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("careerpilot")


# ============================================================
# FASTAPI APP
# ============================================================

app = FastAPI(
    title="CareerPilot AI",
    description="AI-powered career and job recommendation system",
    version="2.0.0",
)


# ============================================================
# CORS CONFIGURATION
# ============================================================
#
# IMPORTANT:
# The frontend is deployed on Render.
#
# Instead of allowing only one exact frontend URL, we allow:
# - localhost development
# - 127.0.0.1 development
# - Render frontend applications
# - HTTPS frontend applications
#
# This prevents the production frontend from being blocked
# because of a small URL/origin difference.
# ============================================================

FRONTEND_URL = os.getenv("FRONTEND_URL", "").strip()

allowed_origins = [
    "http://localhost:5173",
    "http://localhost:5174",
    "http://127.0.0.1:5173",
    "http://127.0.0.1:5174",
]

if FRONTEND_URL:
    allowed_origins.append(FRONTEND_URL.rstrip("/"))

app.add_middleware(
    CORSMiddleware,

    # Local development origins
    allow_origins=allowed_origins,

    # Render / HTTPS production origins
    #
    # This covers:
    # https://careerpilot-ai-frontend-j86v.onrender.com
    # and similar Render frontend URLs.
    allow_origin_regex=(
        r"https://.*\.onrender\.com"
    ),

    # Frontend currently does not need cookies/auth credentials.
    # Keeping this False also makes wildcard-style production
    # CORS behavior safer and simpler.
    allow_credentials=False,

    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)


# ============================================================
# REQUEST MODELS
# ============================================================

class JobAnalysisRequest(BaseModel):
    resume_skills: List[str] = Field(default_factory=list)
    job_description: str


class RoadmapRequest(BaseModel):
    missing_skills: List[str] = Field(default_factory=list)


# ============================================================
# SKILL DICTIONARY
# ============================================================

SKILL_ALIASES = {
    "python": [
        "python",
        "python3",
    ],

    "java": [
        "java",
    ],

    "javascript": [
        "javascript",
        "js",
        "ecmascript",
    ],

    "typescript": [
        "typescript",
        "ts",
    ],

    "c": [
        "c programming",
        "c language",
    ],

    "c++": [
        "c++",
        "cpp",
    ],

    "c#": [
        "c#",
        "c sharp",
        "c-sharp",
    ],

    "sql": [
        "sql",
        "mysql",
        "postgresql",
        "postgres",
        "oracle sql",
    ],

    "mysql": [
        "mysql",
    ],

    "postgresql": [
        "postgresql",
        "postgres",
    ],

    "html": [
        "html",
        "html5",
    ],

    "css": [
        "css",
        "css3",
    ],

    "react": [
        "react",
        "reactjs",
        "react.js",
    ],

    "node.js": [
        "node.js",
        "nodejs",
        "node js",
    ],

    "fastapi": [
        "fastapi",
    ],

    "flask": [
        "flask",
    ],

    "django": [
        "django",
    ],

    "machine learning": [
        "machine learning",
        "machine-learning",
        "machine learning algorithms",
    ],

    "deep learning": [
        "deep learning",
        "deep-learning",
    ],

    "artificial intelligence": [
        "artificial intelligence",
    ],

    "generative ai": [
        "generative ai",
        "genai",
        "gen ai",
    ],

    "nlp": [
        "nlp",
        "natural language processing",
    ],

    "computer vision": [
        "computer vision",
    ],

    "opencv": [
        "opencv",
    ],

    "tensorflow": [
        "tensorflow",
        "tensor flow",
    ],

    "pytorch": [
        "pytorch",
        "py torch",
    ],

    "pandas": [
        "pandas",
    ],

    "numpy": [
        "numpy",
    ],

    "scikit-learn": [
        "scikit-learn",
        "sklearn",
        "scikit learn",
    ],

    "power bi": [
        "power bi",
        "powerbi",
    ],

    "tableau": [
        "tableau",
    ],

    "excel": [
        "excel",
        "microsoft excel",
        "ms excel",
    ],

    "data analysis": [
        "data analysis",
        "data analytics",
        "data analyst",
    ],

    "data visualization": [
        "data visualization",
        "data visualisation",
    ],

    "statistics": [
        "statistics",
        "statistical analysis",
    ],

    "git": [
        "git",
    ],

    "github": [
        "github",
    ],

    "gitlab": [
        "gitlab",
    ],

    "docker": [
        "docker",
    ],

    "aws": [
        "aws",
        "amazon web services",
    ],

    "azure": [
        "azure",
        "microsoft azure",
    ],

    "gcp": [
        "gcp",
        "google cloud",
        "google cloud platform",
    ],

    "rest api": [
        "rest api",
        "restful api",
        "rest apis",
    ],

    "api": [
        "api",
        "apis",
        "application programming interface",
    ],

    "mongodb": [
        "mongodb",
        "mongo db",
    ],

    "sqlite": [
        "sqlite",
    ],

    "linux": [
        "linux",
    ],

    "agile": [
        "agile",
        "scrum",
    ],
}


COMMON_SKILLS = [
    "python",
    "java",
    "javascript",
    "typescript",
    "c",
    "c++",
    "c#",
    "sql",
    "mysql",
    "postgresql",
    "html",
    "css",
    "react",
    "node.js",
    "fastapi",
    "flask",
    "django",
    "machine learning",
    "deep learning",
    "artificial intelligence",
    "generative ai",
    "nlp",
    "computer vision",
    "opencv",
    "tensorflow",
    "pytorch",
    "pandas",
    "numpy",
    "scikit-learn",
    "power bi",
    "tableau",
    "excel",
    "data analysis",
    "data visualization",
    "statistics",
    "git",
    "github",
    "gitlab",
    "docker",
    "aws",
    "azure",
    "gcp",
    "rest api",
    "api",
    "mongodb",
    "sqlite",
    "linux",
    "agile",
]


# ============================================================
# SKILL EXTRACTION
# ============================================================

def _contains_skill(text: str, skill: str) -> bool:
    """
    Detect a skill in text while avoiding many false positives.
    """

    if not text:
        return False

    text_lower = text.lower()

    aliases = SKILL_ALIASES.get(skill, [skill])

    for alias in aliases:
        alias = alias.lower().strip()

        if not alias:
            continue

        # Special handling for C.
        if skill == "c":
            patterns = [
                r"(?<![a-z0-9])c(?![a-z0-9+#])",
                r"\bc programming\b",
                r"\bc language\b",
            ]

            for pattern in patterns:
                if re.search(pattern, text_lower):
                    return True

            continue

        # Escape alias so characters such as + and . are handled.
        escaped = re.escape(alias)

        # Permit punctuation around technologies.
        pattern = (
            r"(?<![a-z0-9])"
            + escaped
            + r"(?![a-z0-9])"
        )

        if re.search(pattern, text_lower):
            return True

    return False


def extract_skills(text: str) -> List[str]:
    """
    Extract supported technical skills from resume/job text.
    """

    if not text:
        return []

    found = []

    for skill in COMMON_SKILLS:
        if _contains_skill(text, skill):
            found.append(skill)

    return found


# ============================================================
# PDF EXTRACTION
# ============================================================

def extract_pdf_text(data: bytes) -> str:
    """
    Extract text from PDF using pypdf.

    pypdf is intentionally imported here so that the rest of
    the backend can still start even if PDF dependency is
    temporarily unavailable.
    """

    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail=(
                "PDF support is not installed on the backend. "
                "Please add 'pypdf' to requirements.txt and redeploy."
            ),
        )

    try:
        reader = PdfReader(io.BytesIO(data))

        pages = []

        for page in reader.pages:
            try:
                page_text = page.extract_text() or ""
                pages.append(page_text)
            except Exception:
                pages.append("")

        return "\n".join(pages)

    except Exception as exc:
        logger.exception("PDF extraction failed")

        raise HTTPException(
            status_code=400,
            detail=f"Could not read PDF: {str(exc)}",
        )


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(data: bytes) -> str:
    """
    Extract text from DOCX files.
    """

    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail=(
                "DOCX support is not installed on the backend. "
                "Please add 'python-docx' to requirements.txt and redeploy."
            ),
        )

    try:
        document = Document(io.BytesIO(data))

        parts = []

        # Paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text:
                parts.append(paragraph.text)

        # Tables
        for table in document.tables:
            for row in table.rows:
                row_text = []

                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)

                if row_text:
                    parts.append(" ".join(row_text))

        return "\n".join(parts)

    except Exception as exc:
        logger.exception("DOCX extraction failed")

        raise HTTPException(
            status_code=400,
            detail=f"Could not read DOCX: {str(exc)}",
        )


# ============================================================
# OLD DOC EXTRACTION
# ============================================================

def extract_doc_text(data: bytes) -> str:
    """
    Basic support for old .doc files.

    Old Microsoft .doc files are OLE binary documents and are
    not reliably supported without additional dependencies.

    The application therefore recommends PDF or DOCX.
    """

    # Try a few decodings as a fallback.
    for encoding in (
        "utf-8",
        "utf-16",
        "latin-1",
    ):
        try:
            text = data.decode(
                encoding,
                errors="ignore",
            )

            # Remove binary/control characters.
            text = re.sub(
                r"[\x00-\x08\x0b\x0c\x0e-\x1f]",
                " ",
                text,
            )

            # Clean whitespace.
            text = re.sub(
                r"\s+",
                " ",
                text,
            ).strip()

            # Accept only if it looks like actual text.
            words = re.findall(
                r"[A-Za-z]{2,}",
                text,
            )

            if len(words) >= 10:
                return text

        except Exception:
            continue

    raise HTTPException(
        status_code=400,
        detail=(
            "Old .doc files are not reliably readable. "
            "Please save the resume as PDF or DOCX and upload again."
        ),
    )


# ============================================================
# GENERAL RESUME READER
# ============================================================

async def read_resume_file(file: UploadFile) -> str:
    """
    Read and extract text from PDF, DOCX or DOC.
    """

    filename = (file.filename or "").strip()

    if not filename:
        raise HTTPException(
            status_code=400,
            detail="No filename provided.",
        )

    filename_lower = filename.lower()

    allowed_extensions = (
        ".pdf",
        ".docx",
        ".doc",
    )

    if not filename_lower.endswith(allowed_extensions):
        raise HTTPException(
            status_code=400,
            detail=(
                "Unsupported file type. "
                "Please upload PDF, DOCX, or DOC."
            ),
        )

    try:
        data = await file.read()
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Could not read uploaded file: {str(exc)}",
        )

    if not data:
        raise HTTPException(
            status_code=400,
            detail="The uploaded file is empty.",
        )

    if filename_lower.endswith(".pdf"):
        return extract_pdf_text(data)

    if filename_lower.endswith(".docx"):
        return extract_docx_text(data)

    return extract_doc_text(data)


# ============================================================
# ROOT ROUTE
# ============================================================

@app.get("/")
def home():
    return {
        "message": "CareerPilot AI backend is running",
        "status": "ok",
        "service": "CareerPilot AI",
        "version": "2.0.0",
        "docs": "/docs",
        "health": "/health",
    }


# ============================================================
# HEALTH ROUTE
# ============================================================

@app.get("/health")
def health():
    return {
        "status": "healthy",
        "service": "CareerPilot AI",
        "version": "2.0.0",
    }


# ============================================================
# CORS TEST ROUTE
# ============================================================

@app.get("/cors-test")
def cors_test():
    return {
        "status": "ok",
        "message": "CORS is configured correctly.",
    }


# ============================================================
# RESUME UPLOAD
# ============================================================

@app.post("/upload-resume")
async def upload_resume(
    file: UploadFile = File(...)
):
    """
    Upload and analyze a resume.

    Frontend request:

        FormData
        key = file

    Response:

        {
            "filename": "...",
            "skills": [...],
            "text_preview": "...",
            "message": "..."
        }
    """

    logger.info(
        "Resume upload received: %s",
        file.filename,
    )

    text = await read_resume_file(file)

    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail=(
                "No readable text was found in the resume. "
                "If this is a scanned PDF, please use a text-based "
                "PDF or DOCX file."
            ),
        )

    skills = extract_skills(text)

    cleaned_preview = re.sub(
        r"\s+",
        " ",
        text,
    ).strip()

    return {
        "filename": file.filename,
        "skills": skills,
        "text_preview": cleaned_preview[:1500],
        "skill_count": len(skills),
        "message": "Resume uploaded and analyzed successfully.",
    }


# ============================================================
# JOB MATCHING
# ============================================================

@app.post("/analyze-job")
def analyze_job(
    request: JobAnalysisRequest
):
    """
    Compare resume skills against job description.
    """

    job_description = (
        request.job_description or ""
    ).strip()

    if not job_description:
        raise HTTPException(
            status_code=400,
            detail="Job description cannot be empty.",
        )

    # Normalize resume skills.
    resume_skills = []

    for skill in request.resume_skills:

        value = str(skill).strip().lower()

        if value and value not in resume_skills:
            resume_skills.append(value)

    # Detect required job skills.
    required_skills = extract_skills(
        job_description
    )

    # No supported skills found.
    if not required_skills:
        return {
            "match_score": 0,
            "matched_skills": [],
            "missing_skills": [],
            "required_skills": [],
            "message": (
                "No supported technical skills were detected "
                "in this job description."
            ),
        }

    matched_skills = []
    missing_skills = []

    for required in required_skills:

        if required in resume_skills:
            matched_skills.append(required)
        else:
            missing_skills.append(required)

    match_score = round(
        (
            len(matched_skills)
            / len(required_skills)
        )
        * 100
    )

    return {
        "match_score": match_score,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "required_skills": required_skills,
        "message": "Job compatibility analyzed successfully.",
    }


# ============================================================
# LEARNING ROADMAP DATA
# ============================================================

ROADMAP_DATA = {

    "python": {
        "level": "Beginner",
        "duration": "2-3 weeks",
        "topics": [
            "Syntax",
            "Functions",
            "OOP",
            "File handling",
            "Libraries",
        ],
        "project": "Build a Resume Skill Analyzer",
    },

    "java": {
        "level": "Beginner",
        "duration": "3-4 weeks",
        "topics": [
            "Syntax",
            "OOP",
            "Collections",
            "Exception handling",
            "JDBC",
        ],
        "project": "Build a Student Management System",
    },

    "sql": {
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            "SELECT",
            "JOINs",
            "GROUP BY",
            "Subqueries",
            "Window functions",
        ],
        "project": "Build a Sales Analytics Database",
    },

    "mysql": {
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            "Tables",
            "Queries",
            "JOINs",
            "Indexes",
            "Database design",
        ],
        "project": "Build a Career Data Management Database",
    },

    "excel": {
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            "Formulas",
            "Functions",
            "Pivot tables",
            "Charts",
            "Lookups",
        ],
        "project": "Build an Employee Analytics Dashboard",
    },

    "power bi": {
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            "Power BI interface",
            "Data cleaning",
            "DAX basics",
            "Charts",
            "Dashboards",
        ],
        "project": "Build a Business Intelligence Dashboard",
    },

    "tableau": {
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            "Tableau interface",
            "Connecting datasets",
            "Charts",
            "Filters",
            "Dashboards",
        ],
        "project": "Build a Sales Dashboard",
    },

    "machine learning": {
        "level": "Intermediate",
        "duration": "3-4 weeks",
        "topics": [
            "Regression",
            "Classification",
            "Clustering",
            "Feature engineering",
            "Evaluation",
        ],
        "project": "Build a Job Salary Prediction Model",
    },

    "deep learning": {
        "level": "Intermediate",
        "duration": "4-5 weeks",
        "topics": [
            "Neural networks",
            "CNNs",
            "Training",
            "Validation",
            "Transfer learning",
        ],
        "project": "Build an Image Classification System",
    },

    "artificial intelligence": {
        "level": "Intermediate",
        "duration": "3-4 weeks",
        "topics": [
            "AI fundamentals",
            "Search",
            "Reasoning",
            "Machine learning",
            "AI applications",
        ],
        "project": "Build an AI Career Recommendation System",
    },

    "generative ai": {
        "level": "Intermediate",
        "duration": "3-4 weeks",
        "topics": [
            "LLMs",
            "Prompt engineering",
            "Embeddings",
            "RAG",
            "Evaluation",
        ],
        "project": "Build a Resume Q&A Assistant",
    },

    "nlp": {
        "level": "Intermediate",
        "duration": "3-4 weeks",
        "topics": [
            "Text preprocessing",
            "Embeddings",
            "Classification",
            "Transformers",
            "Evaluation",
        ],
        "project": "Build a Job Description Classifier",
    },

    "javascript": {
        "level": "Beginner",
        "duration": "2-3 weeks",
        "topics": [
            "ES6",
            "Functions",
            "DOM",
            "Async/Await",
            "APIs",
        ],
        "project": "Build an Interactive Career Dashboard",
    },

    "typescript": {
        "level": "Beginner",
        "duration": "2-3 weeks",
        "topics": [
            "Types",
            "Interfaces",
            "Functions",
            "Generics",
            "Type-safe APIs",
        ],
        "project": "Build a TypeScript Career Dashboard",
    },

    "react": {
        "level": "Beginner",
        "duration": "2-3 weeks",
        "topics": [
            "Components",
            "Props",
            "State",
            "Hooks",
            "API integration",
        ],
        "project": "Build a Job Matching Frontend",
    },

    "html": {
        "level": "Beginner",
        "duration": "1 week",
        "topics": [
            "Semantic HTML",
            "Forms",
            "Tables",
            "Accessibility",
            "Page structure",
        ],
        "project": "Build a Personal Portfolio",
    },

    "css": {
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            "Selectors",
            "Flexbox",
            "Grid",
            "Responsive design",
            "Animations",
        ],
        "project": "Build a Responsive Portfolio",
    },

    "git": {
        "level": "Beginner",
        "duration": "1 week",
        "topics": [
            "Commits",
            "Branches",
            "Merge",
            "Pull requests",
            "GitHub",
        ],
        "project": "Publish CareerPilot AI on GitHub",
    },

    "github": {
        "level": "Beginner",
        "duration": "1 week",
        "topics": [
            "Repositories",
            "Branches",
            "Pull requests",
            "Issues",
            "GitHub Actions",
        ],
        "project": "Deploy a CareerPilot AI project",
    },

    "docker": {
        "level": "Intermediate",
        "duration": "1-2 weeks",
        "topics": [
            "Images",
            "Containers",
            "Dockerfile",
            "Compose",
            "Volumes",
        ],
        "project": "Containerize CareerPilot AI",
    },

    "aws": {
        "level": "Beginner",
        "duration": "2-3 weeks",
        "topics": [
            "EC2",
            "S3",
            "IAM",
            "Networking",
            "Deployment",
        ],
        "project": "Deploy a FastAPI Application",
    },

    "azure": {
        "level": "Beginner",
        "duration": "2-3 weeks",
        "topics": [
            "Azure fundamentals",
            "App Service",
            "Storage",
            "Identity",
            "Deployment",
        ],
        "project": "Deploy a FastAPI application on Azure",
    },

    "opencv": {
        "level": "Intermediate",
        "duration": "2-3 weeks",
        "topics": [
            "Images",
            "Video",
            "Contours",
            "Object detection",
            "Preprocessing",
        ],
        "project": "Build a Traffic Monitoring System",
    },

    "computer vision": {
        "level": "Intermediate",
        "duration": "3-4 weeks",
        "topics": [
            "Image processing",
            "Detection",
            "Classification",
            "YOLO",
            "Evaluation",
        ],
        "project": "Build an Object Detection Application",
    },

    "pandas": {
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            "DataFrames",
            "Series",
            "Data cleaning",
            "Grouping",
            "Data analysis",
        ],
        "project": "Build a Resume Data Analyzer",
    },

    "numpy": {
        "level": "Beginner",
        "duration": "1 week",
        "topics": [
            "Arrays",
            "Indexing",
            "Vectorization",
            "Statistics",
            "Linear algebra",
        ],
        "project": "Build a Data Analysis Tool",
    },

    "fastapi": {
        "level": "Intermediate",
        "duration": "1-2 weeks",
        "topics": [
            "Routes",
            "Request models",
            "File uploads",
            "CORS",
            "Deployment",
        ],
        "project": "Build a Resume Analysis API",
    },

    "mongodb": {
        "level": "Beginner",
        "duration": "2 weeks",
        "topics": [
            "Documents",
            "Collections",
            "CRUD",
            "Queries",
            "Indexes",
        ],
        "project": "Build a Career Profile Database",
    },

    "linux": {
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            "Terminal",
            "Files",
            "Permissions",
            "Processes",
            "Networking",
        ],
        "project": "Deploy CareerPilot AI on Linux",
    },

    "agile": {
        "level": "Beginner",
        "duration": "1 week",
        "topics": [
            "Agile principles",
            "Scrum",
            "Sprints",
            "Backlog",
            "Retrospectives",
        ],
        "project": "Plan CareerPilot AI using Scrum",
    },
}


# ============================================================
# ROADMAP BUILDER
# ============================================================

def build_roadmap_item(skill: str) -> dict:

    skill_key = skill.strip().lower()

    if skill_key in ROADMAP_DATA:

        data = ROADMAP_DATA[skill_key].copy()

        data["skill"] = skill

        return data

    return {
        "skill": skill,
        "level": "Beginner",
        "duration": "1-2 weeks",
        "topics": [
            f"Introduction to {skill}",
            f"Core {skill} concepts",
            "Practical exercises",
            "Real-world applications",
            "Interview preparation",
        ],
        "project": (
            f"Build a practical project using {skill}"
        ),
    }


# ============================================================
# LEARNING ROADMAP ENDPOINT
# ============================================================

@app.post("/learning-roadmap")
def learning_roadmap(
    request: RoadmapRequest
):
    """
    Generate a learning roadmap for missing skills.
    """

    missing = []

    existing_lower = set()

    for skill in request.missing_skills:

        value = str(skill).strip()

        if not value:
            continue

        key = value.lower()

        if key not in existing_lower:
            missing.append(value)
            existing_lower.add(key)

    if not missing:
        raise HTTPException(
            status_code=400,
            detail=(
                "No missing skills supplied. "
                "Analyze a job first."
            ),
        )

    roadmap = [
        build_roadmap_item(skill)
        for skill in missing
    ]

    return {
        "skills": missing,
        "roadmap": roadmap,
        "message": "Learning roadmap generated successfully.",
    }


# ============================================================
# GLOBAL ERROR HANDLER
# ============================================================

@app.get("/api-status")
def api_status():
    """
    Simple endpoint useful for frontend debugging.
    """

    return {
        "backend": "online",
        "service": "CareerPilot AI",
        "cors": "enabled",
        "resume_upload": "/upload-resume",
        "job_matching": "/analyze-job",
        "roadmap": "/learning-roadmap",
    }


# ============================================================
# START SERVER
# ============================================================

if __name__ == "__main__":

    import uvicorn

    # Render provides PORT through environment variables.
    port = int(
        os.environ.get(
            "PORT",
            "8000",
        )
    )

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=port,
    )